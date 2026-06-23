"""Document parser service — h025ai-7.

Extracts text from tender documents and parses structured fields.

Supported formats:
  - DOCX  (python-docx)
  - DOC   (legacy; raw bytes scan + fallback to textract if installed)
  - XLSX  (openpyxl)
  - XLS   (xlrd with formatting_info=False)
  - PDF   (pdfplumber primary; pikepdf fallback for broken PDFs)
  - ZIP   (recursive: re-dispatch inner files)
  - RAR   (rarfile if available)
  - 7Z    (if available)

Two layers:
  1) `extract_text()` — turn a file/blob into a plain UTF-8 string.
  2) `extract_structured()` — same as above + `parse_financial_strict()`
     which pulls НМЦК/guarantee numbers from the text using STRICT MODE
     (only canonical label; never a fallback synonym — see below).

STRICT MODE for financial fields (h025ai-7 / SPEC §8.2):
  We DO NOT use synonyms like "Максимальное значение цены договора" as a
  fallback for НМЦК. That label denotes a рамочный лимит (framework cap),
  not the auction's starting price, and using it produced a 99.6% discount
  bug observed by a Habr-based investigation (research zakupki-html-recon).
  If the canonical label is not found → return None.

ANTI-OUTLIER GUARD:
  When the caller (matcher) sees `discount = 1 - (actual / nmck) > 0.80`,
  it must treat the НМЦК as suspicious and either downgrade the score
  or display a "verify manually" warning. Implemented in matcher.py.

DEPENDENCIES (added to requirements.txt):
  python-docx, openpyxl, xlrd, pdfplumber, pikepdf, rarfile,
  selectolax, beautifulsoup4, lxml
"""
from __future__ import annotations

import hashlib
import io
import logging
import re
import zipfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Iterable

logger = logging.getLogger(__name__)


# ── Public dataclasses ───────────────────────────────────────────


@dataclass
class ParsedDocument:
    """Result of parsing a single document.

    `plain_text` is the canonical full text used for AI extraction.
    `text_by_section` is a best-effort split (useful for source_pages).
    `structured` holds non-text fields like financial block.
    `warnings` lists any anti-outlier / strict-mode flags.
    """

    plain_text: str
    text_length: int
    content_hash: str  # SHA256 hex of file bytes (for caching)
    pages_estimated: int = 0
    text_by_section: dict[str, str] = field(default_factory=dict)
    structured: dict[str, Any] = field(default_factory=dict)
    warnings: list[str] = field(default_factory=list)
    parse_status: str = "parsed"  # 'parsed' | 'partial' | 'failed'
    parse_error: str | None = None

    def to_db_dict(self) -> dict[str, Any]:
        return {
            "plain_text": self.plain_text,
            "text_length": self.text_length,
            "content_hash": self.content_hash,
            "pages_estimated": self.pages_estimated,
            "structured": self.structured,
            "warnings": self.warnings,
            "parse_status": self.parse_status,
            "parse_error": self.parse_error,
        }


# ── Strict mode constants ────────────────────────────────────────

# Canonical label for НМЦК (start/max auction price). Anything else → None.
# Forms observed in real ЕИС documents (lowercased):
NMCCK_CANONICAL_PATTERNS: tuple[str, ...] = (
    r"начальная\s*\(\s*максимальная\s*\)\s*цена\s*контракта",
    r"начальная\s*максимальная\s*цена\s*контракта",
    r"нмцк",  # NMCCK abbreviation
    r"н\(м\)цк",
    r"н\(м\)\s*цк",  # Н(М)ЦК
    r"начальная\s*цена\s*контракта",
    r"начальная\s*цена\s*договора",
)

# Synonyms that MUST NOT be used as NMCCK fallback:
#  - "максимальное значение цены договора" → framework contract limit
#  - "максимальная цена контракта" → also framework; ambiguous
NMCCK_FORBIDDEN_PATTERNS: tuple[str, ...] = (
    r"максимальное\s*значение\s*цены\s*договора",
    r"максимальная\s*цена\s*контракта",
    r"предельная\s*цена",
)

# Guarantee patterns (these are fine to match with fallbacks — less strict)
APP_GUARANTEE_PATTERN = r"обеспечение\s*(?:заявки|участия|заявки\s*на\s*участие)"
CONTRACT_GUARANTEE_PATTERN = r"обеспечение\s*(?:исполнения\s*контракта|контракта|исполнения)"


# ── Helpers ──────────────────────────────────────────────────────


def _sha256(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


def _clean_text(s: str) -> str:
    """Normalize whitespace, fix soft hyphens, strip control chars."""
    if not s:
        return ""
    s = s.replace("\u00ad", "")  # soft hyphen
    s = s.replace("\xa0", " ")  # NBSP
    s = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", " ", s)
    s = re.sub(r"\s+", " ", s)
    return s.strip()


def _detect_extension(filename: str, file_bytes: bytes | None = None) -> str:
    """Return lower-case extension without dot.

    Detection priority:
      1. Filename extension if it's a known type (``docx``/``xlsx`` are ZIP
         internally — extension is authoritative for Office formats)
      2. Magic-byte sniff for unknown / generic extensions (``a.bin`` that
         contains PDF data is treated as PDF)
    """
    KNOWN_EXTS = {"docx", "xlsx", "xls", "doc", "pdf", "zip", "rar", "txt", "rtf"}
    ext = Path(filename).suffix.lower().lstrip(".")
    if ext in KNOWN_EXTS:
        return ext
    if file_bytes:
        if file_bytes[:4] == b"%PDF":
            return "pdf"
        if file_bytes[:4] == b"PK\x03\x04":
            return "zip"
        if file_bytes[:8] == b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1":
            return "xls"  # OLE compound file
        if file_bytes[:2] == b"\xd0\xcf":
            return "doc"
        if file_bytes[:6] in (b"Rar!\x1a\x07", b"Rar!\x1a\x07\x01"):
            return "rar"
    return ext  # fall through to original extension (may be empty)


# ── Format-specific extractors ───────────────────────────────────


def _extract_docx(data: bytes) -> str:
    """Extract text from a .docx (Office Open XML) file."""
    from docx import Document  # python-docx

    doc = Document(io.BytesIO(data))
    parts: list[str] = []
    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            parts.append(t)
    for table in doc.tables:
        for row in table.rows:
            row_txt = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_txt:
                parts.append(row_txt)
    return "\n".join(parts)


def _extract_xlsx(data: bytes) -> str:
    """Extract text from .xlsx (read all sheets, all cells)."""
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(data), data_only=True, read_only=True)
    parts: list[str] = []
    for sheet in wb.worksheets:
        parts.append(f"--- Sheet: {sheet.title} ---")
        for row in sheet.iter_rows(values_only=True):
            line = " | ".join(str(c) for c in row if c is not None)
            if line.strip():
                parts.append(line)
    return "\n".join(parts)


def _extract_xls(data: bytes) -> str:
    """Extract text from legacy .xls (OLE compound file) via xlrd."""
    try:
        import xlrd  # type: ignore
    except ImportError as e:
        raise RuntimeError(
            "xlrd is required to parse .xls files. pip install xlrd==2.0.1"
        ) from e
    book = xlrd.open_workbook(file_contents=data, formatting_info=False)
    parts: list[str] = []
    for sheet in book.sheets():
        parts.append(f"--- Sheet: {sheet.name} ---")
        for r in range(sheet.nrows):
            row = sheet.row_values(r)
            line = " | ".join(str(c) for c in row if c not in (None, ""))
            if line.strip():
                parts.append(line)
    return "\n".join(parts)


def _extract_doc(data: bytes) -> str:
    """Extract text from legacy .doc (OLE compound file).

    Tries several strategies:
      1. antiword via subprocess (if available)
      2. catdoc via subprocess (if available)
      3. Raw bytes scan (last resort — produces garbage for binary streams)
    """
    # Strategy 1: antiword
    for cmd in (["antiword", "-"], ["catdoc", "-"]):
        if _which(cmd[0]):
            try:
                import subprocess

                result = subprocess.run(
                    cmd,
                    input=data,
                    capture_output=True,
                    timeout=30,
                    check=False,
                )
                if result.returncode == 0 and result.stdout:
                    return result.stdout.decode("utf-8", errors="replace")
            except Exception as e:
                logger.debug("doc via %s failed: %s", cmd[0], e)

    # Strategy 2: extract readable ASCII / UTF-16 strings
    # (very lossy but useful as last resort)
    ascii_strings = re.findall(rb"[\x20-\x7e]{6,}", data)
    utf16_strings = re.findall(rb"(?:[\x20-\x7e]\x00){6,}", data)
    parts: list[str] = []
    for s in utf16_strings:
        try:
            parts.append(s.decode("utf-16-le"))
        except UnicodeDecodeError:
            pass
    for s in ascii_strings:
        try:
            parts.append(s.decode("ascii"))
        except UnicodeDecodeError:
            pass
    if parts:
        return "\n".join(parts)
    raise RuntimeError(
        "Could not parse .doc file — install 'antiword' or 'catdoc' for better results"
    )


def _which(name: str) -> str | None:
    import shutil

    return shutil.which(name)


def _extract_pdf(data: bytes) -> tuple[str, int]:
    """Extract text from PDF using pdfplumber.

    Returns (text, page_count). Falls back to pikepdf for broken PDFs.
    """
    text_parts: list[str] = []
    page_count = 0
    try:
        import pdfplumber

        with pdfplumber.open(io.BytesIO(data)) as pdf:
            page_count = len(pdf.pages)
            for i, page in enumerate(pdf.pages, 1):
                try:
                    txt = page.extract_text() or ""
                except Exception as e:
                    logger.warning("pdfplumber failed on page %d: %s", i, e)
                    txt = ""
                if txt:
                    text_parts.append(f"--- Page {i} ---\n{txt}")
    except Exception as e:
        logger.warning("pdfplumber failed, trying pikepdf fallback: %s", e)
        try:
            import pikepdf

            with pikepdf.open(io.BytesIO(data)) as pdf:
                page_count = len(pdf.pages)
                # pikepdf alone doesn't extract text — it's a structure lib
                # Real text extraction still requires pdfplumber.
                text_parts.append(
                    f"--- pikepdf opened PDF with {page_count} pages, but text "
                    f"extraction not available without pdfplumber. Original error: {e} ---"
                )
        except Exception as e2:
            raise RuntimeError(f"Both pdfplumber and pikepdf failed: {e2}") from e2
    return "\n".join(text_parts), page_count


def _extract_zip(data: bytes, original_name: str = "") -> list[tuple[str, bytes]]:
    """Recursively extract supported files from ZIP archives.

    Returns list of (filename, content_bytes). For nested ZIPs, recurses
    one level (we don't go too deep to avoid zip-bombs).
    """
    out: list[tuple[str, bytes]] = []
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as zf:
            # Guard against zip bombs: limit total uncompressed size
            total_uncompressed = sum(i.file_size for i in zf.infolist())
            if total_uncompressed > 500 * 1024 * 1024:  # 500 MB
                raise RuntimeError(
                    f"ZIP too large when uncompressed: {total_uncompressed} bytes"
                )
            for info in zf.infolist():
                if info.is_dir():
                    continue
                # Path traversal protection
                safe_name = Path(info.filename).name
                if safe_name != info.filename and "/" in info.filename:
                    # Strip directory prefix
                    pass
                inner_data = zf.read(info)
                if inner_data[:2] == b"PK" and info.filename.lower().endswith(".zip"):
                    # Nested zip — recurse one level
                    try:
                        out.extend(_extract_zip(inner_data, info.filename))
                    except Exception:
                        out.append((safe_name or info.filename, inner_data))
                else:
                    out.append((safe_name or info.filename, inner_data))
    except zipfile.BadZipFile as e:
        raise RuntimeError(f"Bad ZIP file: {e}") from e
    return out


def _extract_rar(data: bytes) -> list[tuple[str, bytes]]:
    """Extract files from RAR archive (if rarfile/unrar installed)."""
    try:
        import rarfile  # type: ignore
    except ImportError as e:
        raise RuntimeError("rarfile not installed; cannot read .rar") from e
    rf = rarfile.RarFile(io.BytesIO(data))
    out: list[tuple[str, bytes]] = []
    for info in rf.infolist():
        if info.is_dir():
            continue
        out.append((Path(info.filename).name, rf.read(info)))
    return out


# ── Public API ───────────────────────────────────────────────────


def extract_text(filename: str, data: bytes) -> ParsedDocument:
    """Extract plain text from any supported file format.

    Args:
        filename: used for extension detection + record-keeping.
        data: raw file bytes.

    Returns:
        ParsedDocument with text, hash, status.
    """
    content_hash = _sha256(data)
    ext = _detect_extension(filename, data)
    warnings: list[str] = []
    parse_status = "parsed"
    parse_error: str | None = None
    plain = ""
    pages = 0
    text_by_section: dict[str, str] = {}

    try:
        if ext == "docx":
            plain = _extract_docx(data)
        elif ext == "xlsx":
            plain = _extract_xlsx(data)
        elif ext == "xls":
            plain = _extract_xls(data)
        elif ext == "doc":
            plain = _extract_doc(data)
        elif ext == "pdf":
            plain, pages = _extract_pdf(data)
        elif ext == "zip":
            inner = _extract_zip(data, filename)
            sub_docs: list[str] = []
            for inner_name, inner_data in inner:
                inner_ext = _detect_extension(inner_name, inner_data)
                if inner_ext in ("docx", "xlsx", "xls", "doc", "pdf", "txt"):
                    try:
                        sub = extract_text(inner_name, inner_data)
                        sub_docs.append(
                            f"--- File: {inner_name} ---\n{sub.plain_text}"
                        )
                    except Exception as e:
                        warnings.append(f"Failed to parse inner file {inner_name}: {e}")
                elif inner_ext in ("png", "jpg", "jpeg", "tiff", "bmp"):
                    warnings.append(
                        f"Image inside archive {inner_name} requires OCR (not in MVP)"
                    )
                else:
                    warnings.append(f"Skipped unsupported inner file: {inner_name}")
            plain = "\n".join(sub_docs)
        elif ext == "rar":
            inner = _extract_rar(data)
            sub_docs = []
            for inner_name, inner_data in inner:
                try:
                    sub = extract_text(inner_name, inner_data)
                    sub_docs.append(
                        f"--- File: {inner_name} ---\n{sub.plain_text}"
                    )
                except Exception as e:
                    warnings.append(f"Failed to parse RAR entry {inner_name}: {e}")
            plain = "\n".join(sub_docs)
        elif ext in ("txt", "csv", "md"):
            plain = data.decode("utf-8", errors="replace")
        elif ext == "xml":
            # Tender notification XML
            plain = _extract_xml_text(data)
        elif ext in ("png", "jpg", "jpeg", "tiff", "bmp"):
            warnings.append("Image file requires OCR (not in MVP)")
            parse_status = "partial"
        elif ext in ("sig",):
            warnings.append("Signature file — skipped (no content)")
            parse_status = "parsed"
        elif ext == "7z":
            warnings.append("7z archives not supported in MVP")
            parse_status = "partial"
        else:
            warnings.append(f"Unknown file extension '{ext}' — skipped")
            parse_status = "failed"
            parse_error = f"Unknown extension: {ext}"
    except Exception as e:
        logger.exception("Document parsing failed for %s", filename)
        parse_status = "failed"
        parse_error = str(e)[:500]

    plain = _clean_text(plain)
    return ParsedDocument(
        plain_text=plain,
        text_length=len(plain),
        content_hash=content_hash,
        pages_estimated=pages,
        text_by_section=text_by_section,
        structured={},
        warnings=warnings,
        parse_status=parse_status,
        parse_error=parse_error,
    )


def _extract_xml_text(data: bytes) -> str:
    """Extract <text> content from tender notification XML."""
    try:
        from selectolax.parser import XMLParser  # fast C-based XML parser

        tree = XMLParser(data)
        # Collect all text nodes recursively
        texts: list[str] = []
        for node in tree.iter():
            if node.text and node.text.strip():
                texts.append(node.text.strip())
        return "\n".join(texts)
    except ImportError:
        # Fallback to BeautifulSoup
        from bs4 import BeautifulSoup

        soup = BeautifulSoup(data, "lxml")
        return soup.get_text(separator="\n", strip=True)


# ── STRICT MODE financial extractor ──────────────────────────────


def _find_number_near(
    text: str, pattern: str, window: int = 200
) -> tuple[float | None, str | None]:
    """Find first number within `window` chars after a regex match.

    Returns (number, surrounding_quote) or (None, None) if nothing found.
    """
    m = re.search(pattern, text, re.IGNORECASE)
    if not m:
        return None, None
    start = max(0, m.end() - 30)
    end = min(len(text), m.end() + window)
    snippet = text[start:end]
    nums = re.findall(r"(\d{1,3}(?:[\s\u00a0]\d{3})*(?:[.,]\d+)?|\d+[.,]\d+|\d+)", snippet)
    if not nums:
        return None, snippet[:200]
    # Take the first plausible number (>= 1000 RUB to avoid false positives)
    for raw in nums:
        n = _parse_rub_number(raw)
        if n is not None and n >= 1000:
            return n, snippet[:200]
    return None, snippet[:200]


def _parse_rub_number(raw: str) -> float | None:
    """Parse a Russian-formatted number like '12 450 000,00' → 12450000.0."""
    if not raw:
        return None
    s = raw.replace("\u00a0", " ").replace(" ", "")
    if "," in s and "." in s:
        # Assume 12,345,678.90 (US) — keep dot as decimal
        s = s.replace(",", "")
    elif "," in s:
        # Russian: 12 450 000,00 — comma is decimal
        s = s.replace(",", ".")
    try:
        return float(s)
    except ValueError:
        return None


def parse_financial_strict(text: str) -> dict[str, Any]:
    """Extract financial block from text in STRICT MODE.

    Rules:
      1. NMCCK — search ONLY canonical labels. Never fall back to forbidden
         patterns ("максимальное значение цены договора" etc.).
      2. If no canonical NMCCK label found → nmck_rub = None (not 0).
      3. Application/contract guarantees: less strict — common variants OK.
      4. Returns warnings if forbidden patterns are detected (so the
         caller can show "NMCCK ambiguous — verify").

    Returns dict with keys:
      nmck_rub, application_guarantee_rub, application_guarantee_pct,
      contract_guarantee_rub, contract_guarantee_pct,
      nmck_quote, application_guarantee_quote, contract_guarantee_quote,
      warnings
    """
    out: dict[str, Any] = {
        "nmck_rub": None,
        "application_guarantee_rub": None,
        "application_guarantee_pct": None,
        "contract_guarantee_rub": None,
        "contract_guarantee_pct": None,
        "nmck_quote": None,
        "application_guarantee_quote": None,
        "contract_guarantee_quote": None,
        "warnings": [],
    }

    if not text:
        return out

    # Check forbidden patterns — warn but don't use as NMCCK
    for bad in NMCCK_FORBIDDEN_PATTERNS:
        if re.search(bad, text, re.IGNORECASE):
            out["warnings"].append(
                f"Found forbidden NMCCK synonym '{bad}' — ignoring "
                f"(likely рамочный лимит, not auction start price)"
            )

    # NMCCK — canonical only
    for canon in NMCCK_CANONICAL_PATTERNS:
        val, quote = _find_number_near(text, canon)
        if val is not None:
            out["nmck_rub"] = val
            out["nmck_quote"] = quote
            break

    # Application guarantee
    val, quote = _find_number_near(text, APP_GUARANTEE_PATTERN)
    if val is not None:
        out["application_guarantee_rub"] = val
        out["application_guarantee_quote"] = quote
    # Look for percentage near application guarantee
    pct = re.search(
        r"(?:" + APP_GUARANTEE_PATTERN + r").{0,200}?(\d+(?:[.,]\d+)?)\s*%",
        text,
        re.IGNORECASE,
    )
    if pct:
        try:
            out["application_guarantee_pct"] = float(pct.group(1).replace(",", "."))
        except ValueError:
            pass

    # Contract guarantee
    val, quote = _find_number_near(text, CONTRACT_GUARANTEE_PATTERN)
    if val is not None:
        out["contract_guarantee_rub"] = val
        out["contract_guarantee_quote"] = quote
    pct = re.search(
        r"(?:" + CONTRACT_GUARANTEE_PATTERN + r").{0,200}?(\d+(?:[.,]\d+)?)\s*%",
        text,
        re.IGNORECASE,
    )
    if pct:
        try:
            out["contract_guarantee_pct"] = float(pct.group(1).replace(",", "."))
        except ValueError:
            pass

    return out


def apply_anti_outlier_guard(
    nmck: float | None,
    contract_price: float | None,
    max_discount: float = 0.80,
) -> tuple[float | None, str | None]:
    """Apply anti-outlier guard: if discount > max_discount, mark suspicious.

    discount = 1 - (contract_price / nmck)

    Returns (effective_nmck, warning_message_or_none).
    If discount > max_discount (default 80%), the NMCCK is treated as
    suspicious (likely рамочный лимит or parsing error). The effective
    NMCCK returned is None (caller must NOT use it for scoring without
    manual verification).
    """
    if nmck is None or contract_price is None:
        return nmck, None
    if nmck <= 0:
        return nmck, None
    discount = 1 - (contract_price / nmck)
    if discount > max_discount:
        return (
            None,
            f"discount={discount:.2%} > {max_discount:.0%} — NMCCK appears anomalous "
            f"(рамочный лимит or parser error); manual verification required",
        )
    return nmck, None


# ── Module exports ───────────────────────────────────────────────

__all__ = [
    "ParsedDocument",
    "extract_text",
    "parse_financial_strict",
    "apply_anti_outlier_guard",
    "NMCCK_CANONICAL_PATTERNS",
    "NMCCK_FORBIDDEN_PATTERNS",
]