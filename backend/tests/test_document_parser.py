"""Unit tests for document_parser — h025ai-7.

Run with: python -m pytest backend/tests/test_document_parser.py -v

These tests use the real libraries and do NOT mock the parser.
Some tests require optional packages (antiword, rarfile).
"""
import io
import zipfile

import pytest

from app.services.document_parser import (
    NMCCK_CANONICAL_PATTERNS,
    NMCCK_FORBIDDEN_PATTERNS,
    ParsedDocument,
    _detect_extension,
    _parse_rub_number,
    apply_anti_outlier_guard,
    extract_text,
    parse_financial_strict,
)


# ── extract_text ────────────────────────────────────────────────


def test_extract_text_docx_minimal():
    """Build a minimal DOCX in memory and verify text extraction."""
    from docx import Document

    doc = Document()
    doc.add_paragraph("Поставка компьютерного оборудования")
    doc.add_paragraph("Начальная (максимальная) цена контракта: 12 450 000,00 руб.")
    doc.add_paragraph("Обеспечение заявки: 5%")
    doc.add_table(rows=1, cols=2)

    buf = io.BytesIO()
    doc.save(buf)
    data = buf.getvalue()

    result = extract_text("tender.docx", data)
    assert isinstance(result, ParsedDocument)
    assert "компьютерного оборудования" in result.plain_text
    assert "12 450 000" in result.plain_text or "12 450" in result.plain_text
    assert result.parse_status == "parsed"
    assert len(result.content_hash) == 64  # SHA256 hex


def test_extract_text_xlsx_minimal():
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.title = "Смета"
    ws.append(["Позиция", "Количество", "Цена"])
    ws.append(["Ноутбук", 50, 120000])
    ws.append(["Мышь", 50, 1500])
    buf = io.BytesIO()
    wb.save(buf)
    data = buf.getvalue()

    result = extract_text("smeta.xlsx", data)
    assert "Смета" in result.plain_text
    assert "Ноутбук" in result.plain_text
    assert result.parse_status == "parsed"


def test_extract_text_zip_recursive():
    """ZIP containing DOCX should extract text from the inner DOCX."""
    from docx import Document

    doc = Document()
    doc.add_paragraph("Внутренний документ из архива")
    buf = io.BytesIO()
    doc.save(buf)
    inner_docx = buf.getvalue()

    zip_buf = io.BytesIO()
    with zipfile.ZipFile(zip_buf, "w") as zf:
        zf.writestr("readme.txt", "Это README внутри архива".encode("utf-8"))
        zf.writestr("specs/spec.docx", inner_docx)

    result = extract_text("archive.zip", zip_buf.getvalue())
    assert "Внутренний документ" in result.plain_text
    assert "README" in result.plain_text or "readme" in result.plain_text.lower()


def test_extract_text_unknown_extension():
    data = b"some random binary content"
    result = extract_text("weird.zzz", data)
    assert result.parse_status in ("partial", "failed")
    assert result.parse_error is not None


def test_extract_text_txt():
    result = extract_text("readme.txt", "Простой текст\nС новой строки".encode("utf-8"))
    assert "Простой текст" in result.plain_text
    assert "С новой строки" in result.plain_text
    assert result.parse_status == "parsed"


def test_detect_extension_by_magic_bytes():
    # No extension in filename → magic byte sniff
    assert _detect_extension("a", b"%PDF-1.4\n...") == "pdf"
    assert _detect_extension("a", b"PK\x03\x04...") == "zip"


# ── Strict mode financial extraction ────────────────────────────


def test_strict_mode_nmck_canonical_label():
    """Canonical 'Начальная (максимальная) цена контракта' must match."""
    text = """
    Извещение о закупке №123.
    Начальная (максимальная) цена контракта: 12 450 000,00 руб.
    Срок подачи заявок: 28.06.2026.
    """
    fin = parse_financial_strict(text)
    assert fin["nmck_rub"] == 12450000.0
    assert fin["nmck_quote"] is not None
    assert "12 450 000" in fin["nmck_quote"]


def test_strict_mode_nmck_abbreviation():
    """НМЦК abbreviation should also match (it's canonical)."""
    text = "НМЦК: 5 000 000 руб."
    fin = parse_financial_strict(text)
    assert fin["nmck_rub"] == 5000000.0


def test_strict_mode_nmck_forbidden_synonym_returns_none():
    """Forbidden pattern 'Максимальное значение цены договора' must NOT be used."""
    text = """
    Рамочный контракт.
    Максимальное значение цены договора: 1 000 000 000,00 руб.
    """
    fin = parse_financial_strict(text)
    assert fin["nmck_rub"] is None  # strict mode: no fallback
    assert any("forbidden" in w for w in fin["warnings"])


def test_strict_mode_nmck_absent_returns_none():
    """When no NMCCK is present, nmck_rub should be None (not 0!)."""
    text = "Это просто описание тендера без финансовых полей."
    fin = parse_financial_strict(text)
    assert fin["nmck_rub"] is None


def test_strict_mode_guarantees():
    text = """
    Обеспечение заявки: 622 500,00 руб. (5%)
    Обеспечение исполнения контракта: 1 245 000,00 руб. (10%)
    """
    fin = parse_financial_strict(text)
    assert fin["application_guarantee_rub"] == 622500.0
    assert fin["application_guarantee_pct"] == 5.0
    assert fin["contract_guarantee_rub"] == 1245000.0
    assert fin["contract_guarantee_pct"] == 10.0


def test_parse_rub_number():
    assert _parse_rub_number("12 450 000,00") == 12450000.0
    assert _parse_rub_number("12 450 000.00") == 12450000.0
    assert _parse_rub_number("1 000") == 1000.0
    assert _parse_rub_number("invalid") is None


# ── Anti-outlier guard ──────────────────────────────────────────


def test_anti_outlier_guard_normal_discount():
    """Discount <= 80% → NMCCK is OK."""
    nmck, warning = apply_anti_outlier_guard(nmck=1_000_000, contract_price=900_000)
    assert nmck == 1_000_000
    assert warning is None


def test_anti_outlier_guard_extreme_discount():
    """Discount > 80% → NMCCK flagged as anomalous, returns None."""
    # 99.6% discount (the real-world bug from Habr case)
    nmck, warning = apply_anti_outlier_guard(
        nmck=1_000_000_000, contract_price=4_000_000
    )
    assert nmck is None  # caller must not use
    assert warning is not None
    assert "99.60%" in warning or "anomalous" in warning.lower()


def test_anti_outlier_guard_none_inputs():
    """When NMCCK or contract price is None, no warning."""
    nmck, warning = apply_anti_outlier_guard(nmck=None, contract_price=100)
    assert nmck is None
    assert warning is None


def test_canonical_patterns_include_abbreviation():
    """NMCCK_CANONICAL_PATTERNS must include 'нмцк' abbreviation."""
    assert any("нмцк" in p for p in NMCCK_CANONICAL_PATTERNS)


def test_forbidden_patterns_include_framework_limit():
    """NMCCK_FORBIDDEN_PATTERNS must include the framework limit synonym."""
    assert any("максимальное" in p for p in NMCCK_FORBIDDEN_PATTERNS)